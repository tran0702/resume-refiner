
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class CoverLetterBuilder:
    def __init__(self, profile, letter_body):
        self.profile = profile
        self.letter_body = letter_body
        self.doc = Document()
        self.set_margins()

    def set_margins(self):
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

    def add_header(self):
        # Name
        name = self.profile.get('name', 'Your Name')
        heading = self.doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(name)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'

        # Contact Info
        contact_info = []
        if 'phone' in self.profile: contact_info.append(self.profile['phone'])
        if 'email' in self.profile: contact_info.append(self.profile['email'])
        if 'linkedin' in self.profile: contact_info.append(self.profile['linkedin'])
        
        contact_line = self.doc.add_paragraph(" | ".join(contact_info))
        contact_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_line.style.font.size = Pt(10)
        contact_line.style.font.name = 'Times New Roman'
        
        # Add date
        self.doc.add_paragraph().paragraph_format.space_before = Pt(12)
        date_line = self.doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date_line.style.font.name = 'Times New Roman'

    def add_body(self):
        # Determine if we need to split paragraphs
        paragraphs = self.letter_body.split('\n\n')
        for para in paragraphs:
            p = self.doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(12)
            p.style.font.name = 'Times New Roman'
            p.style.font.size = Pt(11)

    def build(self, output_path):
        self.add_header()
        self.add_body()
        self.doc.save(output_path)
        return output_path
