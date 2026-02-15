
import requests
import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

BASE_URL = "http://127.0.0.1:5000/api/upload-resume"

def create_dummy_pdf():
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.drawString(100, 750, "This is a dummy PDF resume.")
    c.drawString(100, 730, "Experience: Software Engineer")
    c.save()
    buffer.seek(0)
    return buffer

def create_dummy_docx():
    buffer = io.BytesIO()
    doc = Document()
    doc.add_paragraph("This is a dummy DOCX resume.")
    doc.add_paragraph("Experience: Python Developer")
    buffer.seek(0)
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def test_upload(file_buffer, filename):
    files = {'file': (filename, file_buffer)}
    try:
        response = requests.post(BASE_URL, files=files)
        if response.status_code == 200:
            print(f"[{filename}] SUCCESS")
            print(f"Extracted: {response.json().get('extracted_text').strip()}")
        else:
            print(f"[{filename}] FAILED: {response.status_code} - {response.text}")
    except Exception as e:
        print(f"[{filename}] ERROR: {str(e)}")

if __name__ == "__main__":
    print("Testing PDF Upload...")
    test_upload(create_dummy_pdf(), "test_resume.pdf")
    
    print("\nTesting DOCX Upload...")
    test_upload(create_dummy_docx(), "test_resume.docx")
