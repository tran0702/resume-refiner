
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ResumeBuilder:
    def __init__(self, master_profile_json):
        """
        master_profile_json: dict containing profile data (experience, education, skills, etc.)
        """
        self.profile = master_profile_json
        self.doc = Document()
        self.set_margins()

    def set_margins(self):
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

    def add_header(self):
        # Name
        name = self.profile.get('name', 'Your Name')
        heading = self.doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(name)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'

        # Contact Info
        contact_info = []
        if 'phone' in self.profile: contact_info.append(self.profile['phone'])
        if 'email' in self.profile: contact_info.append(self.profile['email'])
        if 'linkedin' in self.profile: contact_info.append(self.profile['linkedin'])
        
        contact_line = self.doc.add_paragraph(" | ".join(contact_info))
        contact_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_line.style.font.size = Pt(10)
        contact_line.style.font.name = 'Times New Roman'

    def add_section_title(self, title):
        p = self.doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        # Add bottom border logic (simplified as text underline or dash for now, 
        # real borders require OXML which is complex for this snippet)
        p.add_run("\n" + "_"*90).font.size = Pt(2) # Fake border line appx

    def add_experience(self):
        if 'experience' not in self.profile: return
        self.add_section_title("Experience")
        
        for job in self.profile['experience']:
            # Role + Company + Dates line
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            
            # Use tabs for layout: Role (left) ... Company (right) ?? 
            # Harvard style usually: Company (Left) ... Location (Right) 
            # Next line: Role (Left) ... Dates (Right)
            
            # Simple implementation:
            company = job.get('company', '')
            location = job.get('location', '')
            role = job.get('role', '')
            dates = job.get('dates', '')

            # Line 1: Company (Bold) ... Location
            line1 = p.add_run(f"{company}")
            line1.bold = True
            line1.font.name = 'Times New Roman'
            if location:
                p.add_run(f" — {location}").font.name = 'Times New Roman'

            # Line 2: Role (Italic) ... Dates
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            line2 = p2.add_run(f"{role}")
            line2.italic = True
            line2.font.name = 'Times New Roman'
            if dates:
                p2.add_run(f" ({dates})").font.name = 'Times New Roman'

            # Bullets
            for bullet in job.get('bullets', []):
                bp = self.doc.add_paragraph(style='List Bullet')
                bp.paragraph_format.space_after = Pt(0)
                brun = bp.add_run(bullet)
                brun.font.name = 'Times New Roman'
                brun.font.size = Pt(10.5)

    def add_education(self):
        if 'education' not in self.profile: return
        self.add_section_title("Education")
        
        for edu in self.profile['education']:
            school = edu.get('school', '')
            degree = edu.get('degree', '')
            year = edu.get('year', '')
            
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            
            r1 = p.add_run(school)
            r1.bold = True
            r1.font.name = 'Times New Roman'
            
            if year:
                p.add_run(f", {year}").font.name = 'Times New Roman'
            
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(4)
            r2 = p2.add_run(degree)
            r2.italic = True
            r2.font.name = 'Times New Roman'

    def add_skills(self):
        if 'skills' not in self.profile: return
        self.add_section_title("Skills")
        
        p = self.doc.add_paragraph()
        run = p.add_run(", ".join(self.profile['skills']))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)

    def build(self, output_path):
        self.add_header()
        self.add_education() # Harvard often puts Education first for students, Experience for pros. Let's stick to standard Experience first? Or make it configurable. 
        # Let's do Education first (common in Harvard templates).
        # Actually, let's stick to Experience first if it exists, logic can be refined later.
        # User requested 'Master Profile' refinement, often for professionals. 
        # Let's do: Education, Experience, Skills.
        
        self.add_experience()
        self.add_skills()
        
        self.doc.save(output_path)
        return output_path

if __name__ == "__main__":
    # Test Data
    dummy_profile = {
        "name": "Jane Doe",
        "email": "jane@example.com",
        "phone": "(555) 123-4567",
        "linkedin": "linkedin.com/in/janedoe",
        "education": [
            {"school": "Harvard University", "degree": "B.A. Computer Science", "year": "2020"}
        ],
        "experience": [
            {
                "company": "Tech Corp",
                "role": "Software Engineer",
                "location": "Boston, MA",
                "dates": "2020 - Present",
                "bullets": [
                    "Developed backend APIs using Python and Flask.",
                    "Improved system performance by 30%."
                ]
            }
        ],
        "skills": ["Python", "React", "Docker", "AWS"]
    }
    
    builder = ResumeBuilder(dummy_profile)
    builder.build("backend/templates/docx/test_harvard.docx")
    print("Generated backend/templates/docx/test_harvard.docx")
